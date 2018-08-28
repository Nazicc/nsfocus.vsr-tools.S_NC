#!/usr/bin/env python3
# author: hjxfire
# Copyright (C) 2018 hjxfire <hjxfire@outlook.com>
# https://github.com/hjxfire/nsfocus.vsr-tools.S_NC
# 用法: python3 start.py

import os
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

from openpyxl import load_workbook
import time


HTMLPATH='html/'
WORDPATH='word/'
TEMPLATEPATH= 'template/'
EXCELPATH='excel/'
PDFPATH='pdf/'
PDFOUTPUTPATH='pdfOutput/'

#获取需要处理的文件夹
def getFilesPath(dirPath):
    dirsname=os.listdir(dirPath)
    #删除.DS_Store
    if '.DS_Store' in dirsname:
        dirsname.remove('.DS_Store')
    return dirsname

def processData(dirname,sheet,domainList):
    '''
    第一部分:读取数据
    '''
    #读取
    for file in os.listdir(HTMLPATH+dirname):
        if file.endswith('html'):
            htmlFile=open(HTMLPATH+dirname+'/'+file,'r')
            soup=BeautifulSoup(htmlFile,'html.parser')
            htmlFile.close()
            break

    table_report_table = soup.find_all('table', class_='report_table')
    tr_even = table_report_table[0].find_all('tr', class_='even')

    #=================================================
    # 搜集内容:域名domain
    domain = tr_even[0].td.text.split('/')[2]

    # ================================================
    # 判断是哪种风险,找出为未知风险的报告
    risk = tr_even[1].td.text
    if risk == '未知风险':
        # 从excel中找到相应的域名
        flag = False
        i = 0
        for i in range(len(domainList)):
            if domainList[i].value == domain:
                flag = True
                break
        # 如果没找到,则报错并返回
        if flag == False:
            print(domain + ':汇总表中未找到')
            return 0
        #找到的,则在对应的备注栏填写'站点不可达'
        sheet.cell(i+3,11,'站点不可达')
        return 0
    #================================================
    #搜集内容:时间timeStatistics[]
    for tr in tr_even:
        if tr.find('th',text='时间统计')!=None:
            timeStatistics=tr.td.text.split('\n')[1:4]
            timeStatistics[0]=timeStatistics[0].strip()[3:22]   #timeStart
            timeStatistics[1]=timeStatistics[1].strip()[3:22]   #timeStop
            timeStatistics[2]=timeStatistics[2].strip()[3:]     #timeConsume
            break

    #================================================
    #搜集内容:高中低漏洞数distribut[],漏洞总数sum_all
    sum_all=0
    distribut=tr_even[2].td.text.split('\n')[1:4]
    for i in range(len(distribut)):
        distribut[i]=int(distribut[i].strip()[4:-2])
        sum_all+=distribut[i]
    #================================================
    #搜集内容:检测结果result[],result里面每个元素都是列表,包含4个元素
    table_tmp=[]
    result=[]
    for table in table_report_table:
        if table.find('th',text='漏洞名称')!=None:
            table_tmp.append(table)
    table=table_tmp[-1].find_all('tr')[1:]
    for tr in table:
        td=tr.find_all('td')
        tdClass=td[0].attrs['class']    #这里tdClass是list格式,虽然这里只有1个元素
        if tdClass[0]=='vul-vh':
            result.append([td[0].text.split('\n')[2].strip(),td[1].text.strip(),'高危','EB3323'])
        elif tdClass[0]=='vul-vm':
            result.append([td[0].text.split('\n')[2].strip(),td[1].text.strip(),'中危','F6C143'])
        else:
            result.append([td[0].text.split('\n')[2].strip(),td[1].text.strip(),'低危','4EAC5B'])

    '''
    第二部分:写到word和excel中
    '''
    #从excel中找到相应的域名
    flag=False
    i=0
    for i in range(len(domainList)):
        if domainList[i].value==domain:
            flag=True
            break
    #如果没找到,则报错并返回
    if flag==False:
        print(domain+':汇总表中未找到')
        return 0

    # 获取当前年月日
    timeStr=time.strftime('%Y%m%d',time.localtime())
    #设置文件名
    wordFileName='006'+timeStr+str(i+1).zfill(5)+'a.docx'
    wordFileName_Pdf='006'+timeStr+str(i+1).zfill(5)+'a.pdf'
    pdfFileName='006'+timeStr+str(i+1).zfill(5)+'ad.pdf'
    wordFilePath=WORDPATH+wordFileName
    # 写入excel
    for j in range(len(distribut)):
        #distribut里放的是数字形式,如果为string,则excel里面也会算文字
        sheet.cell(i+3,j+3,distribut[j])
    sheet.cell(i+3,j+4,timeStatistics[0])
    sheet.cell(i+3,j+5,wordFileName_Pdf)
    #从模板word复制到WORDPATH下
    os.system('cp ' + templateFilePath + ' ' + wordFilePath)
    wordFile=Document(wordFilePath)
    tables=wordFile.tables
    #第一个表所有和第二个表的时间
    ##域名
    run=tables[0].cell(0,1).paragraphs[0].add_run(domain)
    run.font.size=Pt(16)
    run.font.bold=1
    ##时间
    for i in range(len(timeStatistics)):
        #1表
        run=tables[0].cell(i+1,1).paragraphs[0].add_run(timeStatistics[i])
        run.font.size = Pt(16)
        run.font.bold = 1
        #2表
        run=tables[1].cell(i+3,1).paragraphs[0].add_run(timeStatistics[i])
        run.font.size = Pt(11)
    #第二个表剩余的内容
    ##域名
    run=tables[1].cell(0,1).paragraphs[0].add_run(domain)
    run.font.size=Pt(11)
    ##漏洞数
    run=tables[1].cell(1,1).paragraphs[0].add_run(str(distribut[0]))
    run.font.size=Pt(11)
    run=tables[1].cell(1,3).paragraphs[0].add_run(str(distribut[1]))
    run.font.size=Pt(11)
    run=tables[1].cell(2,1).paragraphs[0].add_run(str(distribut[2]))
    run.font.size = Pt(11)
    run=tables[1].cell(2,3).paragraphs[0].add_run(str(sum_all))
    run.font.size = Pt(11)
    #第三个表
    for i in range(len(result)):
        new_cells = tables[2].add_row().cells
        #序号
        run=new_cells[0].paragraphs[0].add_run(str(i+1))
        run.font.size=Pt(11)
        run.font.bold=1
        new_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #检测问题
        run=new_cells[1].paragraphs[0].add_run(result[i][0])
        run.font.size=Pt(11)
        #漏洞数
        run=new_cells[2].paragraphs[0].add_run(result[i][1])
        run.font.size = Pt(11)
        new_cells[2].paragraphs[0].alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        #风险程度
        run=new_cells[3].paragraphs[0].add_run(result[i][2])
        run.font.size=Pt(11)
        run.font.bold=1
        run.font.color.rgb=RGBColor.from_string(result[i][3])
        new_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #保存文档
    wordFile.save(wordFilePath)

    '''
    重命名pdf
    '''
    os.system('cp '+PDFPATH+dirname+'/http_'+domain+'.pdf '+PDFOUTPUTPATH+pdfFileName)
    return 0

dirsname=getFilesPath(HTMLPATH)
excelFilePath=EXCELPATH+getFilesPath(EXCELPATH)[0]
templateFilePath= TEMPLATEPATH + getFilesPath(TEMPLATEPATH)[0]
i=0

# 获取工作簿
wb = load_workbook(excelFilePath)
# 获取当前显示的表
sheet = wb.active
domainList = list(sheet.columns)[1][2:]

for dirname in dirsname:
    print(dirname)
    processData(dirname,sheet,domainList)
    i+=1

wb.save(excelFilePath)